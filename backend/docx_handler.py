import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class DocxHandler:
    """Handler for reading and writing DOCX files."""

    def create_docx(self, paragraphs: list, output_path: str):
        """Create a DOCX file from a list of text paragraphs.

        Args:
            paragraphs: List of text strings
            output_path: Path to save the DOCX file
        """
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add paragraphs
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.space_after = Pt(6)
            else:
                doc.add_paragraph("")  # Empty paragraph for spacing

        # Save
        doc.save(output_path)

    def create_docx_from_html(self, html_content: str, output_path: str):
        """Create a DOCX file from HTML content.

        Args:
            html_content: HTML string with <p> tags
            output_path: Path to save the DOCX file
        """
        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Parse HTML for <p> tags
        para_pattern = re.compile(r'<p[^>]*>(.*?)</p>', re.DOTALL)
        matches = para_pattern.findall(html_content)

        if not matches:
            # Fallback: treat entire content as one paragraph
            clean_text = re.sub(r'<[^>]+>', '', html_content).strip()
            if clean_text:
                doc.add_paragraph(clean_text)
        else:
            for match in matches:
                clean_text = re.sub(r'<[^>]+>', '', match).strip()
                if clean_text:
                    p = doc.add_paragraph(clean_text)
                    p.space_after = Pt(6)
                else:
                    doc.add_paragraph("")

        doc.save(output_path)

    def update_docx(self, docx_path: str, html_content: str):
        """Update an existing DOCX file with new content from HTML.

        Args:
            docx_path: Path to existing DOCX file
            html_content: HTML string with updated content
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        # Recreate from HTML
        self.create_docx_from_html(html_content, docx_path)

    def read_docx(self, docx_path: str) -> dict:
        """Read a DOCX file and return its content.

        Args:
            docx_path: Path to DOCX file

        Returns:
            dict with 'paragraphs' (list) and 'content' (HTML string)
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        doc = Document(docx_path)
        paragraphs = []
        html_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                # Create simple HTML representation
                html_parts.append(f"<p>{text}</p>")
            else:
                html_parts.append("<p></p>")

        return {
            "paragraphs": paragraphs,
            "content": "\n".join(html_parts),
        }
